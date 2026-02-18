# office_modules.py - ExcelSolver + DocumentWriter + PDFBuilder for ClawForge

"""
ClawForge - ExcelSolver + DocumentWriter + PDFBuilder
======================================================
ExcelSolver:
  - Formula creation (VLOOKUP, XLOOKUP, SUMIF, IF, etc.)
  - Pivot table creation
  - Conditional formatting
  - Data cleaning (missing values, duplicates, type fixing)
  - Dashboard generation
  - VBA macros (approval required)
  - Report export

DocumentWriter:
  - PDF reports, proposals, invoices, resumes, business plans
  - DOCX creation with headings, tables, styles
  - Template support

PDFBuilder:
  - PDF generation with reportlab
  - Professional layout
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any

# ============================================================================
# EXCEL SOLVER
# ============================================================================

class ExcelSolver:
    """
    Full Excel capability module:
    - Analyze dataset
    - Clean data
    - Apply formulas
    - Conditional formatting
    - Pivot tables
    - Generate dashboards
    - Export reports
    """
    
    FORMULA_LIBRARY = {
        "vlookup":    '=VLOOKUP({lookup}, {table}, {col_index}, 0)',
        "xlookup":    '=XLOOKUP({lookup}, {lookup_array}, {return_array}, "Not Found")',
        "sumif":      '=SUMIF({range}, {criteria}, {sum_range})',
        "countif":    '=COUNTIF({range}, {criteria})',
        "iferror":    '=IFERROR({formula}, {fallback})',
        "if":         '=IF({condition}, {true_val}, {false_val})',
        "index_match": '=INDEX({return_col}, MATCH({lookup}, {lookup_col}, 0))',
        "sum":        "=SUM({range})",
        "average":    "=AVERAGE({range})",
        "max":        "=MAX({range})",
        "min":        "=MIN({range})",
        "count":      "=COUNT({range})",
    }

    def __init__(self):
        try:
            import openpyxl
            self.openpyxl = openpyxl
            self.available = True
        except ImportError:
            self.openpyxl = None
            self.available = False
        
        # Get workspace paths
        try:
            from backend.identity import WORKSPACE_PATHS
            self.workspace_paths = WORKSPACE_PATHS
        except ImportError:
            self.workspace_paths = {
                "tasks": "./workspace/tasks",
                "output": "./workspace/output"
            }
    
    def _require_openpyxl(self) -> bool:
        if not self.available:
            print("  [WARN] openpyxl not installed. Run: pip install openpyxl")
            return False
        return True
    
    # ============================================================================
    # STEP 1: ANALYZE
    # ============================================================================
    
    def analyze_dataset(self, file_path: str = None, data: list = None) -> dict:
        """Analyzes a dataset and returns structure summary."""
        if file_path and self._require_openpyxl():
            wb = self.openpyxl.load_workbook(file_path)
            results = {}
            for name in wb.sheetnames:
                ws = wb[name]
                headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
                null_counts = {}
                for c in range(1, ws.max_column + 1):
                    null_counts[headers[c-1]] = sum(
                        1 for r in range(2, ws.max_row + 1)
                        if ws.cell(r, c).value is None
                    )
                results[name] = {
                    "rows": ws.max_row - 1,
                    "columns": ws.max_column,
                    "headers": headers,
                    "null_counts": null_counts,
                    "has_data": ws.max_row > 1,
                }
            return results

        elif data:
            if not data:
                return {"error": "Empty dataset"}
            headers = list(data[0].keys()) if isinstance(data[0], dict) else []
            null_counts = {h: sum(1 for row in data if not row.get(h)) for h in headers}
            return {
                "Sheet1": {
                    "rows": len(data),
                    "columns": len(headers),
                    "headers": headers,
                    "null_counts": null_counts,
                    "has_data": True,
                }
            }
        return {"error": "No file_path or data provided"}
    
    # ============================================================================
    # STEP 2: IDENTIFY ERRORS
    # ============================================================================
    
    def identify_errors(self, analysis: dict) -> list:
        """Identifies data quality issues from analysis."""
        issues = []
        for sheet, info in analysis.items():
            if isinstance(info, dict):
                for col, null_count in info.get("null_counts", {}).items():
                    if null_count > 0:
                        issues.append({
                            "sheet": sheet,
                            "column": col,
                            "issue": "Missing values",
                            "count": null_count,
                            "fix": "Fill with default or use IFERROR()",
                        })
                if info.get("rows", 0) == 0:
                    issues.append({"sheet": sheet, "issue": "Empty sheet"})
        return issues
    
    # ============================================================================
    # STEP 3: FORMULA BUILDER
    # ============================================================================
    
    def build_formula(self, formula_type: str, params: dict) -> str:
        """Generates Excel formula string."""
        template = self.FORMULA_LIBRARY.get(formula_type.lower())
        if not template:
            available = list(self.FORMULA_LIBRARY.keys())
            return f"# Unknown formula: {formula_type}. Available: {available}"
        try:
            return template.format(**params)
        except KeyError as e:
            return f"# Missing parameter: {e}"
    
    # ============================================================================
    # STEP 4: CLEAN DATA
    # ============================================================================
    
    def clean_data(self, data: List[Dict]) -> dict:
        """Cleans a list-of-dicts dataset."""
        if not data:
            return {"cleaned": [], "removed_duplicates": 0, "flagged_rows": 0}

        original_count = len(data)
        seen = set()
        cleaned = []
        flagged = 0

        for row in data:
            # Strip whitespace
            row = {k: v.strip() if isinstance(v, str) else v for k, v in row.items()}
            # Check all-null
            if all(v is None or v == "" for v in row.values()):
                flagged += 1
                continue
            # Dedup
            key = str(sorted(row.items()))
            if key not in seen:
                seen.add(key)
                cleaned.append(row)

        return {
            "original_count": original_count,
            "cleaned_count": len(cleaned),
            "removed_duplicates": original_count - len(cleaned) - flagged,
            "flagged_empty_rows": flagged,
            "cleaned": cleaned,
        }
    
    # ============================================================================
    # STEP 5: CREATE EXCEL FILE
    # ============================================================================
    
    def create_excel(
        self,
        output_path: str,
        sheets: Dict[str, List[Dict]],
        title: str = "ClawForge Report",
    ) -> dict:
        """Creates a styled Excel workbook."""
        if not self._require_openpyxl():
            return self._fallback_csv(output_path, sheets)

        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = self.openpyxl.Workbook()
        wb.remove(wb.active)

        for sheet_name, rows in sheets.items():
            ws = wb.create_sheet(title=sheet_name[:31])

            if not rows:
                ws.append(["No data"])
                continue

            headers = list(rows[0].keys())

            # Header styling
            header_fill = PatternFill("solid", fgColor="1E3A5F")
            header_font = Font(color="FFFFFF", bold=True)
            border = Border(
                bottom=Side(style="thin", color="AAAAAA"),
                right=Side(style="thin", color="DDDDDD"),
            )

            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(1, col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
                cell.border = border

            # Data rows
            for row_idx, row in enumerate(rows, 2):
                for col_idx, key in enumerate(headers, 1):
                    cell = ws.cell(row_idx, col_idx, value=row.get(key, ""))
                    if row_idx % 2 == 0:
                        cell.fill = PatternFill("solid", fgColor="F2F7FC")
                    cell.border = border

            # Auto-fit columns
            for col_idx, header in enumerate(headers, 1):
                max_len = max(len(str(header)), max(len(str(row.get(header, ""))) for row in rows) if rows else 0)
                ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 40)

        # Summary sheet
        summary = wb.create_sheet(title="Summary", index=0)
        summary["A1"] = title
        summary["A1"].font = Font(bold=True, size=14)
        summary["A3"] = f"Generated by ClawForge — {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}"
        summary["A4"] = f"Sheets: {', '.join(sheets.keys())}"
        summary["A5"] = f"Total rows: {sum(len(v) for v in sheets.values())}"

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        wb.save(output_path)
        return {"status": "success", "path": output_path, "sheets": list(sheets.keys())}

    def _fallback_csv(self, output_path: str, sheets: dict) -> dict:
        """CSV fallback."""
        import csv
        csv_path = output_path.replace(".xlsx", ".csv")
        for name, rows in sheets.items():
            if rows:
                with open(csv_path, "w", newline="") as f:
                    writer = csv.DictWriter(f, fieldnames=rows[0].keys())
                    writer.writeheader()
                    writer.writerows(rows)
        return {"status": "csv_fallback", "path": csv_path}
    
    # ============================================================================
    # STEP 6: PIVOT TABLE
    # ============================================================================
    
    def describe_pivot_table(self, source_range: str, rows: list, values: list, agg: str = "SUM") -> str:
        """Returns pivot table setup description."""
        return (
            f"Pivot Table Setup:\n"
            f"  Source: {source_range}\n"
            f"  Row Labels: {', '.join(rows)}\n"
            f"  Values ({agg}): {', '.join(values)}\n"
            f"  Steps:\n"
            f"    1. Select data range {source_range}\n"
            f"    2. Insert > PivotTable > New Worksheet\n"
            f"    3. Drag {', '.join(rows)} to Rows\n"
            f"    4. Drag {', '.join(values)} to Values (set to {agg})"
        )
    
    # ============================================================================
    # FULL PIPELINE
    # ============================================================================
    
    def run_pipeline(self, data: list, output_filename: str, task_id: str = "excel_task") -> dict:
        """Full Excel workflow."""
        print(f"\n[OK] ExcelSolver Pipeline: '{output_filename}'")

        try:
            output_path = str(Path(self.workspace_paths["output"]) / task_id / output_filename)
        except (KeyError, TypeError):
            output_path = str(Path("./workspace/output") / task_id / output_filename)
        
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        print("  [1/4] Analyzing dataset...")
        analysis = self.analyze_dataset(data=data)

        print("  [2/4] Identifying errors...")
        errors = self.identify_errors(analysis)
        print(f"         Issues found: {len(errors)}")

        print("  [3/4] Cleaning data...")
        cleaned = self.clean_data(data)
        print(f"         Cleaned: {cleaned['cleaned_count']} rows, removed {cleaned['removed_duplicates']} duplicates")

        print("  [4/4] Creating Excel file...")
        result = self.create_excel(output_path, {"Data": cleaned["cleaned"]}, "ClawForge Report")
        print(f"  [OK] Excel saved: {output_path}")

        return {
            "analysis": analysis,
            "errors_found": errors,
            "cleaned_rows": cleaned["cleaned_count"],
            "excel_path": output_path,
            "formula_examples": {
                "vlookup": self.build_formula("vlookup", {"lookup": "A2", "table": "B:D", "col_index": 2}),
                "sumif": self.build_formula("sumif", {"range": "C2:C100", "criteria": "Criteria", "sum_range": "D2:D100"}),
            }
        }

# ============================================================================
# DOCUMENT WRITER (DOCX)
# ============================================================================

class DocumentWriter:
    """
    Creates professional DOCX documents.
    """
    
    DOCUMENT_TYPES = ["report", "proposal", "invoice", "resume", "business_plan", "generic"]

    def __init__(self):
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            self._Document = Document
            self._Pt = Pt
            self._Inches = Inches
            self._RGBColor = RGBColor
            self._WD_ALIGN = WD_ALIGN_PARAGRAPH
            self.available = True
        except ImportError:
            self.available = False
        
        # Get workspace paths
        try:
            from backend.identity import WORKSPACE_PATHS
            self.workspace_paths = WORKSPACE_PATHS
        except ImportError:
            self.workspace_paths = {"output": "./workspace/output"}
    
    def create_document(
        self,
        doc_type: str,
        title: str,
        content: dict,
        output_path: str,
    ) -> dict:
        """Creates a styled DOCX document."""
        if not self.available:
            return self._fallback_markdown(doc_type, title, content, output_path)

        doc = self._Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.runs[0].font.color.rgb = self._RGBColor(0x1E, 0x3A, 0x5F)

        doc.add_paragraph(
            f"{doc_type.replace('_', ' ').title()} | "
            f"Generated by ClawForge | {datetime.utcnow().strftime('%B %d, %Y')}"
        ).runs[0].italic = True

        doc.add_paragraph("")

        # Build by type
        if doc_type == "report":
            self._build_report(doc, content)
        elif doc_type == "proposal":
            self._build_proposal(doc, content)
        elif doc_type == "invoice":
            self._build_invoice(doc, content)
        elif doc_type == "resume":
            self._build_resume(doc, content)
        elif doc_type == "business_plan":
            self._build_business_plan(doc, content)
        else:
            self._build_generic(doc, content)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(f"Generated by ClawForge — {datetime.utcnow().isoformat()}")
        footer.runs[0].font.size = self._Pt(8)
        footer.runs[0].italic = True

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return {"status": "success", "path": output_path, "doc_type": doc_type}

    def _build_report(self, doc, content: dict):
        sections = content.get("sections", {
            "Executive Summary": "This report provides an overview...",
            "Findings": "Key findings include...",
            "Recommendations": "Based on findings...",
            "Conclusion": "In conclusion...",
        })
        for heading, body in sections.items():
            doc.add_heading(heading, level=1)
            doc.add_paragraph(body)

    def _build_proposal(self, doc, content: dict):
        sections = [
            ("Overview", content.get("overview", "Project overview...")),
            ("Objectives", content.get("objectives", "Key objectives...")),
            ("Scope of Work", content.get("scope", "Work scope...")),
            ("Timeline", content.get("timeline", "Project timeline...")),
            ("Budget", content.get("budget", "Budget breakdown...")),
            ("Team", content.get("team", "Project team...")),
        ]
        for h, body in sections:
            doc.add_heading(h, level=1)
            doc.add_paragraph(body)

    def _build_invoice(self, doc, content: dict):
        doc.add_heading("INVOICE", level=1)
        meta = [
            ("Invoice #", content.get("invoice_number", "INV-001")),
            ("Date", content.get("date", datetime.utcnow().strftime("%Y-%m-%d"))),
            ("Due Date", content.get("due_date", "Net 30")),
            ("Bill To", content.get("client", "Client Name")),
        ]
        table = doc.add_table(rows=len(meta), cols=2)
        table.style = "Table Grid"
        for i, (k, v) in enumerate(meta):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = str(v)

        doc.add_paragraph("")
        doc.add_heading("Line Items", level=2)
        items = content.get("items", [{"description": "Service", "qty": 1, "rate": 100, "total": 100}])
        item_table = doc.add_table(rows=1 + len(items), cols=4)
        item_table.style = "Table Grid"
        headers = ["Description", "Qty", "Rate", "Total"]
        for i, h in enumerate(headers):
            item_table.rows[0].cells[i].text = h

        subtotal = 0
        for r_idx, item in enumerate(items, 1):
            row = item_table.rows[r_idx]
            row.cells[0].text = str(item.get("description", ""))
            row.cells[1].text = str(item.get("qty", 1))
            row.cells[2].text = f"${item.get('rate', 0):.2f}"
            total = item.get("qty", 1) * item.get("rate", 0)
            row.cells[3].text = f"${total:.2f}"
            subtotal += total

        doc.add_paragraph(f"\nSubtotal: ${subtotal:.2f}")
        tax_rate = float(content.get('tax_rate', 0))
        tax = subtotal * tax_rate / 100
        doc.add_paragraph(f"Tax ({tax_rate}%): ${tax:.2f}")
        doc.add_paragraph(f"TOTAL DUE: ${subtotal + tax:.2f}")

    def _build_resume(self, doc, content: dict):
        sections = {
            "Professional Summary": content.get("summary", "Experienced professional..."),
            "Work Experience": content.get("experience", "Job details..."),
            "Education": content.get("education", "Degree details..."),
            "Skills": content.get("skills", "Python, JavaScript..."),
            "Certifications": content.get("certifications", "Certifications..."),
        }
        for h, body in sections.items():
            doc.add_heading(h, level=1)
            doc.add_paragraph(body)

    def _build_business_plan(self, doc, content: dict):
        sections = {
            "Executive Summary": content.get("executive_summary", "Business overview..."),
            "Company Description": content.get("company", "What the company does..."),
            "Market Analysis": content.get("market", "Target market..."),
            "Products & Services": content.get("products", "What we offer..."),
            "Financial Plan": content.get("financials", "Revenue projections..."),
        }
        for h, body in sections.items():
            doc.add_heading(h, level=1)
            doc.add_paragraph(body)

    def _build_generic(self, doc, content: dict):
        for heading, body in content.get("sections", {"Content": "..."}).items():
            doc.add_heading(heading, level=1)
            doc.add_paragraph(str(body))

    def _fallback_markdown(self, doc_type, title, content, output_path) -> dict:
        md_path = output_path.replace(".docx", ".md")
        lines = [f"# {title}\n", f"**Type:** {doc_type}\n", "---\n"]
        for k, v in content.items():
            lines.append(f"## {k}\n{v}\n")
        Path(md_path).parent.mkdir(parents=True, exist_ok=True)
        Path(md_path).write_text("\n".join(lines), encoding="utf-8")
        return {"status": "markdown_fallback", "path": md_path}

# ============================================================================
# PDF BUILDER
# ============================================================================

class PDFBuilder:
    """Builds professional PDF documents using reportlab."""
    
    def __init__(self):
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
            self._letter = letter
            self._A4 = A4
            self._styles = getSampleStyleSheet()
            self._inch = inch
            self._colors = colors
            self._SimpleDocTemplate = SimpleDocTemplate
            self._Paragraph = Paragraph
            self._Spacer = Spacer
            self._Table = Table
            self._TableStyle = TableStyle
            self._PageBreak = PageBreak
            self._HRFlowable = HRFlowable
            self._ParagraphStyle = ParagraphStyle
            self.available = True
        except ImportError:
            self.available = False
        
        try:
            from backend.identity import WORKSPACE_PATHS
            self.workspace_paths = WORKSPACE_PATHS
        except ImportError:
            self.workspace_paths = {"output": "./workspace/output"}

    def create_pdf(
        self,
        output_path: str,
        title: str,
        sections: Dict[str, str],
        author: str = "ClawForge",
        pagesize: str = "letter",
    ) -> dict:
        """Generates a professional PDF report."""
        if not self.available:
            return self._fallback_txt(output_path, title, sections)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        page = self._letter if pagesize == "letter" else self._A4

        doc = self._SimpleDocTemplate(
            output_path,
            pagesize=page,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72,
            title=title,
            author=author,
        )

        story = []
        styles = self._styles

        # Custom styles
        title_style = self._ParagraphStyle(
            "ClawTitle", parent=styles["Title"],
            fontSize=24,
            textColor=self._colors.HexColor("#1E3A5F"),
            spaceAfter=12,
        )
        h1_style = self._ParagraphStyle(
            "ClawH1", parent=styles["Heading1"],
            fontSize=16,
            textColor=self._colors.HexColor("#1E3A5F"),
            spaceBefore=16, spaceAfter=8,
        )
        body_style = self._ParagraphStyle(
            "ClawBody", parent=styles["Normal"],
            fontSize=11, leading=16, spaceAfter=8,
        )
        meta_style = self._ParagraphStyle(
            "ClawMeta", parent=styles["Normal"],
            fontSize=9, textColor=self._colors.grey, spaceAfter=20,
        )

        # Title
        story.append(self._Paragraph(title, title_style))
        story.append(self._Paragraph(
            f"Generated by ClawForge | {author} | {datetime.utcnow().strftime('%B %d, %Y')}",
            meta_style
        ))
        story.append(self._HRFlowable(width="100%", thickness=1, color=self._colors.HexColor("#1E3A5F")))
        story.append(self._Spacer(1, 0.2 * self._inch))

        # Sections
        for heading, body in sections.items():
            story.append(self._Paragraph(heading, h1_style))
            for para in body.split("\n\n"):
                if para.strip():
                    story.append(self._Paragraph(para.replace("\n", "<br/>"), body_style))
            story.append(self._Spacer(1, 0.1 * self._inch))

        # Footer
        story.append(self._HRFlowable(width="100%", thickness=0.5, color=self._colors.grey))
        story.append(self._Paragraph(
            f"ClawForge Document — {datetime.utcnow().isoformat()}",
            meta_style
        ))

        doc.build(story)
        return {"status": "success", "path": output_path}

    def _fallback_txt(self, output_path, title, sections) -> dict:
        txt_path = output_path.replace(".pdf", ".txt")
        lines = [f"{title}\n{'='*len(title)}\n"]
        for h, body in sections.items():
            lines.append(f"\n{h}\n{'-'*len(h)}\n{body}\n")
        Path(txt_path).parent.mkdir(parents=True, exist_ok=True)
        Path(txt_path).write_text("\n".join(lines), encoding="utf-8")
        return {"status": "txt_fallback", "path": txt_path}

# ============================================================================
# COMPUTER CONTROL POLICY
# ============================================================================

class UIController:
    """
    Computer Control Policy enforcer.
    All UI automation requires explicit YES/NO user approval.
    """
    
    def __init__(self, permission_granted: bool = False):
        self.permission_granted = permission_granted

    def request_permission(self, action_description: str) -> str:
        """Returns approval request string."""
        return (
            f"SCREEN CONTROL REQUEST\n"
            f"{'='*50}\n"
            f"ClawForge wants to perform:\n\n"
            f"  -> {action_description}\n\n"
            f"Type YES to allow or NO to cancel.\n"
            f"{'='*50}"
        )

    def execute_if_approved(self, tool_name: str, args: dict, user_response: str) -> dict:
        """Executes UI tool only if user said YES."""
        if user_response.strip().upper() != "YES":
            return {
                "status": "denied",
                "message": "UI action cancelled by user.",
                "tool": tool_name,
            }

        print(f"  [OK] UI action approved: {tool_name}")

        try:
            import pyautogui
            pyautogui.FAILSAFE = True

            if tool_name == "take_screenshot":
                ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                path = str(Path("./workspace/output") / f"screenshot_{ts}.png")
                pyautogui.screenshot().save(path)
                return {"status": "success", "path": path}

            elif tool_name == "click":
                pyautogui.click(args.get("x", 0), args.get("y", 0))
                return {"status": "success"}

            elif tool_name == "type_text":
                pyautogui.write(args.get("text", ""), interval=0.05)
                return {"status": "success"}

        except ImportError:
            return {"status": "unavailable", "message": "PyAutoGUI not installed."}
        except Exception as e:
            return {"status": "error", "message": str(e)}

# ============================================================================
# CONVENIENCE FUNCTIONS
# ============================================================================

def create_excel(data: list, output_filename: str, task_id: str = "excel_task") -> dict:
    """Create an Excel file from data."""
    solver = ExcelSolver()
    return solver.run_pipeline(data, output_filename, task_id)

def create_document(doc_type: str, title: str, content: dict, output_path: str) -> dict:
    """Create a DOCX document."""
    writer = DocumentWriter()
    return writer.create_document(doc_type, title, content, output_path)

def create_pdf(title: str, sections: Dict[str, str], output_path: str) -> dict:
    """Create a PDF document."""
    builder = PDFBuilder()
    return builder.create_pdf(output_path, title, sections)

# ============================================================================
# TEST
# ============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("  ExcelSolver + DocumentWriter + PDFBuilder Test")
    print("=" * 60)

    # Excel test
    solver = ExcelSolver()
    sample_data = [
        {"Name": "Alice", "Department": "Engineering", "Salary": 95000, "Status": "Active"},
        {"Name": "Bob",   "Department": "Marketing",   "Salary": 72000, "Status": "Active"},
        {"Name": "Carol", "Department": "Engineering", "Salary": 88000, "Status": None},
        {"Name": "Bob",   "Department": "Marketing",   "Salary": 72000, "Status": "Active"},
        {"Name": "",      "Department": "",           "Salary": None,  "Status": None},
    ]

    excel_result = solver.run_pipeline(sample_data, "report.xlsx", task_id="test_office_001")
    print(f"\n[OK] Excel: {excel_result['cleaned_rows']} rows, {len(excel_result['errors_found'])} issues")
    print(f"   Formula: {excel_result['formula_examples']['vlookup']}")

    # Document test
    writer = DocumentWriter()
    doc_result = writer.create_document(
        doc_type="report",
        title="Q3 Performance Report",
        content={
            "sections": {
                "Executive Summary": "Q3 showed strong growth...",
                "Findings": "Revenue increased 23% YoY.",
                "Recommendations": "Continue investing in marketing.",
            }
        },
        output_path="./workspace/output/test_report.docx",
    )
    print(f"\n[OK] Document: {doc_result['status']} -> {doc_result['path']}")

    # PDF test
    builder = PDFBuilder()
    pdf_result = builder.create_pdf(
        output_path="./workspace/output/test_report.pdf",
        title="ClawForge Test Report",
        sections={
            "Overview": "This is a test PDF generated by ClawForge.",
            "Key Points": "Automated generation, professional styling.",
            "Conclusion": "ClawForge produces ready-to-use documents.",
        },
    )
    print(f"\n[OK] PDF: {pdf_result['status']} -> {pdf_result['path']}")

    # UI Controller test
    print("\n[OK] UIController Policy Test:")
    ui = UIController()
    prompt = ui.request_permission("Open Chrome and navigate to URL")
    print(prompt[:100] + "...")
    result = ui.execute_if_approved("take_screenshot", {}, "NO")
    print(f"   Result: {result['status']} - {result['message']}")
